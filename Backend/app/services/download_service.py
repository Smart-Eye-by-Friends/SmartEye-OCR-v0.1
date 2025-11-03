"""
Project Download Service
========================

데이터베이스에 저장된 분석 결과를 활용하여 프로젝트 단위의 통합 텍스트 및
Word 문서를 생성합니다. CombinedResult 테이블을 캐시로 사용하며, 각 페이지의
최신(TextVersion.is_current=True) 자동 포맷팅 텍스트를 모아 제공합니다.
"""

from __future__ import annotations

import io
from datetime import datetime
from typing import Dict, List, Optional, Tuple

from loguru import logger
from sqlalchemy.orm import Session, selectinload

from ..models import CombinedResult, Page, Project, TextVersion

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:  # pragma: no cover - optional dependency
    Document = None  # type: ignore[assignment]
    WD_ALIGN_PARAGRAPH = None  # type: ignore[assignment]
    logger.error("python-docx 라이브러리가 설치되지 않았습니다. pip install python-docx")


# -----------------------------------------------------------------------------
# 헬퍼 함수
# -----------------------------------------------------------------------------

def _fetch_project_with_pages(db: Session, project_id: int) -> Project:
    project = (
        db.query(Project)
        .options(selectinload(Project.pages))
        .filter(Project.project_id == project_id)
        .one_or_none()
    )
    if not project:
        raise ValueError(f"프로젝트 ID {project_id}를 찾을 수 없습니다.")
    return project


def _fetch_current_text_versions(db: Session, page_ids: List[int]) -> Dict[int, TextVersion]:
    if not page_ids:
        return {}
    versions = (
        db.query(TextVersion)
        .filter(
            TextVersion.page_id.in_(page_ids),
            TextVersion.is_current.is_(True),
        )
        .all()
    )
    return {version.page_id: version for version in versions}


def _latest_text_timestamp(versions: Dict[int, TextVersion]) -> Optional[datetime]:
    timestamps = [version.created_at for version in versions.values() if version.created_at]
    return max(timestamps) if timestamps else None


def _combined_result_is_fresh(
    combined_result: CombinedResult,
    latest_text_time: Optional[datetime],
) -> bool:
    if latest_text_time is None:
        return combined_result.combined_text is not None
    if combined_result.updated_at is None:
        return False
    return combined_result.updated_at >= latest_text_time


def _format_combined_sections(pages: List[Page], versions: Dict[int, TextVersion]) -> Tuple[str, Dict[str, int]]:
    sections: List[str] = []
    total_words = 0
    total_characters = 0

    for page in sorted(pages, key=lambda p: p.page_number):
        version = versions.get(page.page_id)
        header = f"─── 페이지 {page.page_number}"
        if version:
            header += f" (Version: {version.version_number} - {version.version_type}) ───"
            content = version.content or ""
            total_words += len(content.split())
            total_characters += len(content)
        else:
            header += " (내용 없음) ───"
            content = ""
        sections.append(f"{header}\n\n{content}".rstrip())

    combined_text = "\n\n".join(sections)
    stats = {
        "total_pages": len(pages),
        "total_words": total_words,
        "total_characters": total_characters,
    }
    return combined_text, stats


def _upsert_combined_result(
    db: Session,
    project_id: int,
    combined_text: str,
    stats: Dict[str, int],
) -> CombinedResult:
    record = (
        db.query(CombinedResult)
        .filter(CombinedResult.project_id == project_id)
        .one_or_none()
    )
    now = datetime.utcnow()
    if record:
        record.combined_text = combined_text
        record.combined_stats = stats
        record.updated_at = now
    else:
        record = CombinedResult(
            project_id=project_id,
            combined_text=combined_text,
            combined_stats=stats,
            generated_at=now,
            updated_at=now,
        )
        db.add(record)
    db.commit()
    db.refresh(record)
    return record


# -----------------------------------------------------------------------------
# 공개 API
# -----------------------------------------------------------------------------

def generate_combined_text(
    db: Session,
    project_id: int,
    *,
    use_cache: bool = True,
) -> Dict[str, object]:
    """
    프로젝트의 최신 텍스트 버전을 통합하여 반환합니다.
    CombinedResult 테이블을 캐시로 사용합니다.
    """
    project = _fetch_project_with_pages(db, project_id)
    page_ids = [page.page_id for page in project.pages]
    versions = _fetch_current_text_versions(db, page_ids)
    latest_version_time = _latest_text_timestamp(versions)

    combined_record = (
        db.query(CombinedResult)
        .filter(CombinedResult.project_id == project_id)
        .one_or_none()
    )

    if use_cache and combined_record and _combined_result_is_fresh(combined_record, latest_version_time):
        logger.info("CombinedResult 캐시 사용: project_id=%s", project_id)
        stats = combined_record.combined_stats or {}
        generated_at = combined_record.updated_at or combined_record.generated_at or datetime.utcnow()
        return {
            "project_id": project_id,
            "project_name": project.project_name,
            "combined_text": combined_record.combined_text or "",
            "stats": stats,
            "generated_at": generated_at,
        }

    combined_text, stats = _format_combined_sections(project.pages, versions)
    combined_record = _upsert_combined_result(db, project_id, combined_text, stats)

    return {
        "project_id": project_id,
        "project_name": project.project_name,
        "combined_text": combined_text,
        "stats": stats,
        "generated_at": combined_record.updated_at or combined_record.generated_at or datetime.utcnow(),
    }


def generate_word_document(
    db: Session,
    project_id: int,
    *,
    use_cache: bool = True,
) -> Tuple[str, io.BytesIO]:
    """
    통합 텍스트를 기반으로 Word(.docx) 문서를 생성하고 파일 이름과 스트림을 반환합니다.
    """
    if Document is None or WD_ALIGN_PARAGRAPH is None:
        raise ImportError("python-docx 라이브러리가 필요합니다. pip install python-docx")

    combined_data = generate_combined_text(db, project_id, use_cache=use_cache)
    project_name = combined_data.get("project_name") or f"프로젝트 {project_id}"
    combined_text = combined_data.get("combined_text", "")

    document = Document()
    title = document.add_heading(project_name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_paragraph = document.add_paragraph(
        f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    meta_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    stats = combined_data.get("stats", {})
    total_pages = stats.get("total_pages", 0)
    stats_paragraph = document.add_paragraph(f"총 페이지: {total_pages}개")
    stats_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.add_paragraph("─" * 60)

    sections = [segment.strip() for segment in combined_text.split("─── 페이지 ") if segment.strip()]
    for index, section in enumerate(sections):
        lines = section.split("\n")
        header = lines[0]
        content_lines = lines[1:]

        document.add_heading(f"페이지 {header.split()[0]}", level=2)
        for paragraph in content_lines:
            paragraph = paragraph.strip()
            if paragraph:
                document.add_paragraph(paragraph)
        if index < len(sections) - 1:
            document.add_page_break()

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    filename = f"SmartEyeSsen_{project_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return filename, file_stream


def generate_document(
    db: Session,
    project_id: int,
    *,
    output: str = "text",
    use_cache: bool = True,
):
    """
    통합 텍스트 또는 Word 문서를 생성합니다.
    output="text"  -> dict 반환 (combined_text 등)
    output="docx" -> (filename, BytesIO) 반환
    """
    if output == "docx":
        return generate_word_document(db, project_id, use_cache=use_cache)
    return generate_combined_text(db, project_id, use_cache=use_cache)


__all__ = [
    "generate_document",
    "generate_combined_text",
    "generate_word_document",
]
