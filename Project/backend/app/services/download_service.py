# -*- coding: utf-8 -*-
"""
SmartEyeSsen 통합 다운로드 서비스 (Phase 3.3)
=========================================

프로젝트의 모든 페이지 텍스트를 통합하고 Word(.docx) 문서로 생성합니다.
Mock DB (`mock_pages`, `mock_text_versions`)를 사용합니다.
"""

from typing import List, Dict, Optional, Tuple, Any
from loguru import logger
from datetime import datetime
import io

# python-docx 임포트 (requirements.txt에 추가 필요)
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    logger.error("python-docx 라이브러리가 설치되지 않았습니다. pip install python-docx")
    Document = None # 오류 방지용

# Mock DB 함수 임포트 (batch_analysis.py로부터)
try:
    from .batch_analysis import (
        get_project_mock,
        get_pages_for_project_mock,
        get_latest_version_mock, # 텍스트 버전 조회 함수
        mock_text_versions, # ✅ 성능 최적화를 위한 직접 임포트
        # ✅ Phase 3.3 캐싱 함수 임포트
        save_combined_result_mock,
        get_combined_result_mock,
        is_cache_valid_mock
    )
except ImportError:
    logger.error("Mock DB 함수를 batch_analysis.py에서 임포트할 수 없습니다.")
    # 임시 함수 정의
    def get_project_mock(project_id: int) -> Optional[Dict]: return None
    def get_pages_for_project_mock(project_id: int) -> List[Dict]: return []
    def get_latest_version_mock(page_id: int) -> Optional[Dict]: return None
    mock_text_versions = [] # 빈 리스트로 초기화
    def save_combined_result_mock(project_id: int, combined_text: str, stats: Dict) -> None: pass
    def get_combined_result_mock(project_id: int) -> Optional[Dict]: return None
    def is_cache_valid_mock(project_id: int) -> bool: return False

# ============================================================================
# 통합 텍스트 생성 서비스 함수
# ============================================================================

def generate_combined_text(project_id: int) -> Dict[str, Any]:
    """
    프로젝트의 모든 페이지에 대한 최신 텍스트 버전을 통합하여 반환합니다.
    ✅ combined_results 캐싱 로직 적용
    """
    logger.info(f"서비스: 통합 텍스트 생성 요청 - ProjectID={project_id}")

    # ✅ 1단계: 캐시 확인 및 유효성 검사
    if is_cache_valid_mock(project_id):
        cached_result = get_combined_result_mock(project_id)
        if cached_result:
            logger.info(f"서비스: 캐시에서 통합 텍스트 반환 - ProjectID={project_id}")
            return {
                "project_id": project_id,
                "combined_text": cached_result['combined_text'],
                "stats": cached_result['stats'],
                "generated_at": cached_result['generated_at'].isoformat()
            }

    # ✅ 2단계: 캐시 미스 또는 무효 - 새로 생성
    logger.info(f"서비스: 캐시 미스 또는 무효 - 새로 통합 텍스트 생성 - ProjectID={project_id}")

    project = get_project_mock(project_id)
    if not project:
        raise ValueError(f"프로젝트 ID {project_id} 없음")

    # 페이지 번호 순서대로 페이지 목록 조회
    pages = get_pages_for_project_mock(project_id)
    if not pages:
        logger.warning(f"서비스: 프로젝트 ID {project_id}에 페이지가 없습니다.")
        return {
            "project_id": project_id,
            "combined_text": "",
            "stats": {"total_pages": 0, "total_words": 0, "total_characters": 0},
            "generated_at": datetime.now().isoformat()
        }

    combined_sections: List[str] = []
    total_words = 0
    total_characters = 0

    # ✅ 성능 최적화: 모든 페이지의 버전을 한 번에 조회 후 딕셔너리로 변환 (O(n²) → O(n))
    page_ids = [p['page_id'] for p in pages]
    versions_dict = {
        v['page_id']: v for v in mock_text_versions
        if v['page_id'] in page_ids and v['is_current']
    }

    for page in pages:
        page_id = page['page_id']
        page_number = page['page_number']

        # 딕셔너리에서 O(1) 조회
        latest_version = versions_dict.get(page_id)

        if latest_version and latest_version.get('content'):
            content = latest_version['content']
            # 페이지 구분선 추가
            section = f"─── 페이지 {page_number} (Version: {latest_version['version_number']} - {latest_version['version_type']}) ───\n\n"
            section += content
            combined_sections.append(section)

            # 통계 계산
            total_words += len(content.split())
            total_characters += len(content)
        else:
            logger.warning(f"서비스: 페이지 ID {page_id} (번호: {page_number})의 최신 텍스트 버전을 찾을 수 없습니다.")
            combined_sections.append(f"─── 페이지 {page_number} (내용 없음) ───")

    # 전체 텍스트 조합
    combined_text = "\n\n".join(combined_sections)

    stats = {
        "total_pages": len(pages),
        "total_words": total_words,
        "total_characters": total_characters
    }

    # ✅ 3단계: 새로 생성한 결과를 캐시에 저장
    save_combined_result_mock(project_id, combined_text, stats)
    logger.info(f"서비스: 통합 텍스트 생성 완료 및 캐시 저장 - ProjectID={project_id}, Pages={stats['total_pages']}")

    return {
        "project_id": project_id,
        "combined_text": combined_text,
        "stats": stats,
        "generated_at": datetime.now().isoformat()
    }

# ============================================================================
# Word 문서 생성 서비스 함수
# ============================================================================

def generate_word_document(project_id: int) -> Tuple[str, io.BytesIO]:
    """
    프로젝트의 통합 텍스트를 사용하여 Word(.docx) 문서를 생성하고 파일 스트림을 반환합니다.
    (Phase 3.3 계획서 의사코드 기반)
    """
    if not Document:
        raise ImportError("python-docx 라이브러리가 필요합니다.")

    logger.info(f"서비스: Word 문서 생성 요청 - ProjectID={project_id}")

    # 1. 통합 텍스트 생성 (캐시 우선 로직 필요)
    combined_data = generate_combined_text(project_id)
    combined_text = combined_data["combined_text"]
    project = get_project_mock(project_id)
    if not project:
         raise ValueError(f"프로젝트 ID {project_id} 없음") # generate_combined_text에서 이미 체크됨

    # 2. python-docx 문서 생성 시작
    doc = Document()

    # 기본 글꼴 설정 (예: 나눔고딕) - 시스템에 폰트 설치 필요
    # style = doc.styles['Normal']
    # style.font.name = 'NanumGothic'
    # style.font.size = Pt(12)

    # 2-1. 문서 제목
    title = doc.add_heading(project.get("project_name", f"프로젝트 {project_id}"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2-2. 메타데이터
    doc.add_paragraph(f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}", style='Caption').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(f"총 페이지: {combined_data['stats']['total_pages']}개", style='Caption').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph("─" * 70) # 구분선

    # 2-3. 페이지별 내용 추가
    # 페이지 구분선을 기준으로 텍스트 분리
    page_sections = combined_text.split("─── 페이지 ")
    for i, section in enumerate(page_sections):
        if not section.strip() or i == 0: # 첫 번째 분리 결과는 비어있을 수 있음
            continue

        lines = section.strip().split("\n")
        page_header = lines[0].strip() # 예: "1 (Version: 2 - auto_formatted) ───"
        page_content = "\n".join(lines[1:]).strip()

        # 페이지 제목 (Heading 2 스타일 적용)
        try:
             page_number_str = page_header.split(" ")[0]
             doc.add_heading(f"페이지 {page_number_str}", level=2)
        except:
             doc.add_heading(f"페이지 {i}", level=2) # 페이지 번호 추출 실패 시 순서 사용

        # 페이지 내용 (본문)
        # 여러 빈 줄을 하나로 줄이고, 각 문단을 별도 단락으로 추가
        paragraphs = [p.strip() for p in page_content.split("\n") if p.strip()]
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph)

        # 페이지 나누기 (마지막 페이지 제외)
        if i < len(page_sections) -1 :
             doc.add_page_break()


    # 3. 메모리 스트림에 저장
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0) # 스트림 포인터를 처음으로 이동

    # 4. 파일 이름 생성 및 스트림 반환
    filename = f"SmartEyeSsen_{project_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    logger.info(f"서비스: Word 문서 생성 완료 - FileName='{filename}'")

    return filename, file_stream