# -*- coding: utf-8 -*-
"""
SmartEyeSsen 학습지 분석 API 서버
기존 8_4worksheet_analysis_gradio_notebook.py의 핵심 파이프라인을 FastAPI로 변환
"""

import os
import sys
import cv2
import json
import time
import base64
import io
import colorsys
import random
from collections import Counter
from typing import Optional
from PIL import Image, ImageDraw, ImageFont
import numpy as np

# FastAPI 및 관련 패키지
from fastapi import FastAPI, File, Form, UploadFile, HTTPException
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
import uvicorn

# AI 및 OCR 관련 패키지  
import torch
from huggingface_hub import hf_hub_download
import pytesseract
import openai
from loguru import logger
import platform
# VARCO Vision OCR을 위한 추가 import
from transformers import AutoProcessor, LlavaOnevisionForConditionalGeneration
import re

# 워드 문서 생성을 위한 패키지
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Windows에서 Tesseract 경로 설정
if platform.system() == "Windows":
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# 로그 설정
logger.remove()
logger.add(sys.stderr, level="INFO")

# 디바이스 설정
device = 'cuda:0' if torch.cuda.is_available() else 'cpu'
print(f"✅ 사용 디바이스: {device}")

# FastAPI 앱 생성
app = FastAPI(
    title="SmartEyeSsen 학습지 분석 API",
    description="시각 장애 아동을 위한 AI 기반 학습지 분석 및 텍스트 변환 시스템",
    version="1.0.0"
)

# CORS 설정 (Vue.js 프론트엔드에서 접근할 수 있도록)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # 프로덕션에서는 특정 도메인만 허용
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 정적 파일 제공 (결과 이미지 및 JSON 파일)
os.makedirs("static", exist_ok=True)
app.mount("/static", StaticFiles(directory="static"), name="static")


class VarcoVisionOCR:
    """VARCO Vision 2.0 OCR 처리 클래스"""
    
    def __init__(self):
        self.model = None
        self.processor = None
        self.device = device
        self.model_name = "NCSOFT/VARCO-VISION-2.0-1.7B-OCR"
        
    def initialize(self):
        """모델 초기화 (필요할 때만 로드)"""
        if self.model is None:
            try:
                logger.info("VARCO Vision OCR 모델 로드 중...")
                self.model = LlavaOnevisionForConditionalGeneration.from_pretrained(
                    self.model_name,
                    torch_dtype=torch.float16,
                    attn_implementation="sdpa",
                    device_map="auto",
                )
                self.processor = AutoProcessor.from_pretrained(self.model_name)
                logger.info("VARCO Vision OCR 모델 로드 완료")
                return True
            except Exception as e:
                logger.error(f"VARCO Vision OCR 모델 로드 실패: {e}")
                return False
        return True
    
    def preprocess_image(self, image):
        """OCR 성능 향상을 위한 이미지 전처리"""
        # OpenCV BGR을 PIL RGB로 변환
        if len(image.shape) == 3 and image.shape[2] == 3:
            image_pil = Image.fromarray(cv2.cvtColor(image, cv2.COLOR_BGR2RGB))
        else:
            image_pil = Image.fromarray(image)
        
        # 이미지 크기 조정 (OCR 성능 향상)
        w, h = image_pil.size
        target_size = 2304
        if max(w, h) < target_size:
            scaling_factor = target_size / max(w, h)
            new_w = int(w * scaling_factor)
            new_h = int(h * scaling_factor)
            image_pil = image_pil.resize((new_w, new_h))
            logger.info(f"이미지 크기 조정: {w}x{h} -> {new_w}x{new_h}")
        
        return image_pil
    
    def extract_text(self, image_crop):
        """VARCO Vision을 사용한 텍스트 추출"""
        if not self.initialize():
            return ""
        
        try:
            # 이미지 전처리
            image_pil = self.preprocess_image(image_crop)
            
            # 대화 형식으로 요청 구성
            conversation = [
                {
                    "role": "user",
                    "content": [
                        {"type": "image", "image": image_pil},
                        {"type": "text", "text": "<ocr>"},
                    ],
                },
            ]
            
            # 입력 텐서 생성
            inputs = self.processor.apply_chat_template(
                conversation,
                add_generation_prompt=True,
                tokenize=True,
                return_dict=True,
                return_tensors="pt"
            ).to(self.model.device, torch.float16)
            
            # 텍스트 생성
            with torch.no_grad():
                generate_ids = self.model.generate(
                    **inputs, 
                    max_new_tokens=1024,
                    temperature=0.1,  # 일관성 있는 결과를 위해 낮은 온도
                    do_sample=False   # 결정론적 결과
                )
            
            # 결과 디코딩
            generate_ids_trimmed = [
                out_ids[len(in_ids):] for in_ids, out_ids in zip(inputs.input_ids, generate_ids)
            ]
            output = self.processor.decode(generate_ids_trimmed[0], skip_special_tokens=True)
            
            # 텍스트 정리
            cleaned_text = self.clean_ocr_output(output)
            return cleaned_text
            
        except Exception as e:
            logger.error(f"VARCO Vision OCR 실패: {e}")
            return ""
    
    def clean_ocr_output(self, text):
        """OCR 출력 텍스트 정리"""
        if not text:
            return ""
        
        # 불필요한 특수 문자나 토큰 제거
        text = text.strip()
        
        # 여러 줄바꿈을 단일 줄바꿈으로 정리
        text = re.sub(r'\n\s*\n', '\n', text)
        
        return text


class WorksheetAnalyzer:
    """학습지 분석기 클래스 - Gradio 버전에서 이식"""
    
    def __init__(self):
        self.model = None
        self.device = device
        self.layout_info = []
        self.ocr_results = []
        self.api_results = []
        # 🆕 VARCO Vision OCR 추가
        self.varco_ocr = VarcoVisionOCR()
        self.use_varco_ocr = True  # VARCO OCR 사용 여부

    def download_model(self, model_choice="SmartEyeSsen"):
        """사전 훈련된 DocLayout-YOLO 모델 다운로드"""
        models = {
            "doclaynet_docsynth": {
                "repo_id": "juliozhao/DocLayout-YOLO-DocLayNet-Docsynth300K_pretrained",
                "filename": "doclayout_yolo_doclaynet_imgsz1120_docsynth_pretrain.pt"
            },
            "docstructbench": {
                "repo_id": "juliozhao/DocLayout-YOLO-DocStructBench",
                "filename": "doclayout_yolo_docstructbench_imgsz1024.pt"
            },
            "docsynth300k": {
                "repo_id": "juliozhao/DocLayout-YOLO-DocSynth300K-pretrain",
                "filename": "doclayout_yolo_docsynth300k_imgsz1600.pt"
            },
            "SmartEyeSsen": {
                "repo_id": "AkJeond/SmartEyeSsen",
                "filename": "best_tuned_model.pt"
            }
        }

        selected_model = models.get(model_choice, models["SmartEyeSsen"])

        try:
            logger.info(f"모델 다운로드 중: {selected_model['repo_id']}")
            filepath = hf_hub_download(
                repo_id=selected_model["repo_id"],
                filename=selected_model["filename"]
            )
            logger.info(f"모델 다운로드 완료: {filepath}")
            return filepath
        except Exception as e:
            logger.error(f"모델 다운로드 실패: {e}")
            raise

    def load_model(self, model_path):
        """DocLayout-YOLO 모델 로드"""
        try:
            # DocLayout-YOLO 임포트 (여기서 임포트하여 모델이 필요할 때만 로드)
            try:
                from doclayout_yolo import YOLOv10
            except ImportError:
                logger.error("DocLayout-YOLO가 설치되지 않았습니다. 설치 가이드를 확인하세요.")
                return False

            logger.info("모델 로드 중...")
            self.model = YOLOv10(model_path, task='predict')
            self.model.to(self.device)
            if hasattr(self.model, 'training'):
                self.model.training = False
            logger.info("모델 로드 완료!")
            return True
        except Exception as e:
            logger.error(f"모델 로드 실패: {e}")
            return False

    def analyze_layout(self, image, model_choice="SmartEyeSsen"):
        """레이아웃 분석"""
        try:
            logger.info("레이아웃 분석 시작...")

            # 임시 이미지 저장
            temp_path = "temp_image.jpg"
            cv2.imwrite(temp_path, image)

            # 모델 설정
            if model_choice == "SmartEyeSsen":
                imgsz = 1024
                conf = 0.25
            elif model_choice == "docsynth300k":
                imgsz = 1600
                conf = 0.15
            else:  # doclaynet_docsynth
                imgsz = 1024
                conf = 0.25

            # 분석 실행
            results = self.model.predict(
                temp_path,
                imgsz=imgsz,
                conf=conf,
                iou=0.45,
                device=self.device
            )

            # 결과 추출
            boxes = results[0].boxes.xyxy.cpu().numpy()
            classes = results[0].boxes.cls.cpu().numpy()
            confs = results[0].boxes.conf.cpu().numpy()
            class_names = self.model.names

            layout_info = []
            for i, (box, cls, conf) in enumerate(zip(boxes, classes, confs)):
                x1, y1, x2, y2 = map(int, box)
                cls_id = int(cls)

                try:
                    cls_name = class_names[cls_id]
                except IndexError:
                    cls_name = f"unknown_{cls_id}"

                area = (x2 - x1) * (y2 - y1)
                if area < 100:  # 너무 작은 영역 제외
                    continue

                layout_info.append({
                    'id': i,
                    'class_name': cls_name,
                    'confidence': float(conf),
                    'box': [int(x1), int(y1), int(x2), int(y2)],
                    'width': int(x2 - x1),
                    'height': int(y2 - y1),
                    'area': area
                })

            self.layout_info = layout_info
            logger.info(f"레이아웃 분석 완료: {len(layout_info)}개 영역 감지")
            return layout_info

        except Exception as e:
            logger.error(f"레이아웃 분석 실패: {e}")
            return []

    def perform_ocr(self, image, use_varco=None):
        """OCR 처리 - VARCO Vision 또는 Tesseract 선택 가능"""
        target_classes = [
            'title', 'plain_text', 'abandon_text',
            'table_caption', 'table_footnote',
            'isolated_formula', 'formula_caption', 'question_type',
            'question_text', 'question_number', 'list'
        ]

        # OCR 엔진 결정
        if use_varco is None:
            use_varco = self.use_varco_ocr

        ocr_results = []
        
        logger.info(f"OCR 처리 시작... 엔진: {'VARCO Vision' if use_varco else 'Tesseract'}")
        logger.info(f"총 {len(self.layout_info)}개 레이아웃 요소 중 OCR 대상 필터링")
        
        # VARCO OCR 초기화 (필요한 경우)
        if use_varco:
            if not self.varco_ocr.initialize():
                logger.warning("VARCO Vision 초기화 실패, Tesseract으로 폴백")
                use_varco = False

        target_count = 0

        for layout in self.layout_info:
            cls_name = layout['class_name'].lower()
            
            if cls_name not in target_classes:
                continue
                
            target_count += 1
            logger.info(f"OCR 대상 {target_count}: ID {layout['id']} - 클래스 '{cls_name}'")

            x1, y1, x2, y2 = layout['box']
            x1 = max(0, x1)
            y1 = max(0, y1)
            x2 = min(image.shape[1], x2)
            y2 = min(image.shape[0], y2)

            cropped_img = image[y1:y2, x1:x2]

            try:
                if use_varco:
                    # 🆕 VARCO Vision OCR 사용
                    text = self.varco_ocr.extract_text(cropped_img)
                else:
                    # 기존 Tesseract OCR 사용
                    pil_img = Image.fromarray(cropped_img)
                    custom_config = r'--oem 3 --psm 6'
                    text = pytesseract.image_to_string(
                        pil_img,
                        lang='kor+eng',
                        config=custom_config
                    ).strip()

                if len(text) > 1:
                    ocr_results.append({
                        'id': layout['id'],
                        'class_name': cls_name,
                        'coordinates': [x1, y1, x2, y2],
                        'text': text,
                        'ocr_engine': 'VARCO Vision' if use_varco else 'Tesseract'  # 🆕 엔진 정보 추가
                    })
                    logger.info(f"✅ OCR 성공: ID {layout['id']} ({cls_name}) - '{text[:50]}...' ({len(text)}자)")
                else:
                    logger.warning(f"⚠️ OCR 결과 없음: ID {layout['id']} ({cls_name})")

            except Exception as e:
                logger.error(f"OCR 실패: ID {layout['id']} - {e}")
                
                # VARCO 실패시 Tesseract으로 폴백
                if use_varco:
                    try:
                        logger.info(f"VARCO 실패, Tesseract으로 재시도: ID {layout['id']}")
                        pil_img = Image.fromarray(cropped_img)
                        custom_config = r'--oem 3 --psm 6'
                        text = pytesseract.image_to_string(
                            pil_img,
                            lang='kor+eng',
                            config=custom_config
                        ).strip()
                        
                        if len(text) > 1:
                            ocr_results.append({
                                'id': layout['id'],
                                'class_name': cls_name,
                                'coordinates': [x1, y1, x2, y2],
                                'text': text,
                                'ocr_engine': 'Tesseract (fallback)'
                            })
                            logger.info(f"✅ Tesseract 폴백 성공: ID {layout['id']}")
                    except Exception as fallback_error:
                        logger.error(f"Tesseract 폴백도 실패: ID {layout['id']} - {fallback_error}")

        self.ocr_results = ocr_results
        logger.info(f"OCR 처리 완료: {len(ocr_results)}개 텍스트 블록")
        return ocr_results

    def call_openai_api(self, image, api_key):
        """OpenAI Vision API 호출"""
        if not api_key:
            logger.warning("API 키가 제공되지 않아 AI 설명을 건너뜁니다.")
            return []

        target_classes = ['figure', 'table']
        api_results = []

        try:
            client = openai.OpenAI(api_key=api_key)
            logger.info("OpenAI API 처리 시작...")
        except Exception as e:
            logger.error(f"OpenAI 클라이언트 초기화 실패: {e}")
            return []

        prompts = {
            'figure': "이 그림(figure)의 내용을 간단히 요약해 주세요.",
            'table': "이 표(table)의 주요 내용을 요약해 주세요."
        }

        system_prompt = """당신은 시각 장애 아동을 위한 학습 AI 비서입니다.
시각 자료의 내용을 한국어로 간결하고 명확하게 설명해주세요.
설명은 음성으로 변환될 수 있도록 직접적이고 이해하기 쉽게 작성해주세요."""

        for layout in self.layout_info:
            cls_name = layout['class_name'].lower()
            if cls_name not in target_classes:
                continue

            x1, y1, x2, y2 = layout['box']
            cropped_img = image[y1:y2, x1:x2]

            # 이미지를 base64로 인코딩
            pil_img = Image.fromarray(cv2.cvtColor(cropped_img, cv2.COLOR_BGR2RGB))
            buffered = io.BytesIO()
            pil_img.save(buffered, format="PNG")
            img_base64 = base64.b64encode(buffered.getvalue()).decode("utf-8")

            prompt = prompts.get(cls_name, f"이 {cls_name}의 내용을 간단히 설명해 주세요.")

            try:
                response = client.chat.completions.create(
                    model="gpt-4-turbo",
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {
                            "role": "user",
                            "content": [
                                {"type": "text", "text": prompt},
                                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_base64}"}}
                            ]
                        }
                    ],
                    temperature=0.2,
                    max_tokens=600
                )

                description = response.choices[0].message.content.strip()

                api_results.append({
                    'id': layout['id'],
                    'class_name': cls_name,
                    'coordinates': [x1, y1, x2, y2],
                    'description': description
                })

                logger.info(f"API 응답 완료: ID {layout['id']} - {cls_name}")

            except Exception as e:
                logger.error(f"API 요청 실패: ID {layout['id']} - {e}")

        self.api_results = api_results
        logger.info(f"OpenAI API 처리 완료: {len(api_results)}개 설명 생성")
        return api_results

    def visualize_results(self, image):
        """결과 시각화"""
        img_result = image.copy()
        overlay = image.copy()

        # 클래스별 색상 생성
        random.seed(42)

        unique_classes = list(set(layout['class_name'] for layout in self.layout_info))
        class_colors = {}

        for i, cls_name in enumerate(unique_classes):
            h = i / max(1, len(unique_classes))
            s = 0.8
            v = 0.9
            r, g, b = colorsys.hsv_to_rgb(h, s, v)
            class_colors[cls_name] = (int(b * 255), int(g * 255), int(r * 255))

        # 바운딩 박스 그리기
        for layout in self.layout_info:
            x1, y1, x2, y2 = layout['box']
            cls_name = layout['class_name']
            color = class_colors[cls_name]

            # 반투명 오버레이
            cv2.rectangle(overlay, (x1, y1), (x2, y2), color, -1)

            # 테두리
            cv2.rectangle(img_result, (x1, y1), (x2, y2), color, 2)

            # 라벨
            label = f"{cls_name} ({layout['confidence']:.2f})"
            labelSize, _ = cv2.getTextSize(label, cv2.FONT_HERSHEY_SIMPLEX, 0.5, 1)
            y1_label = max(y1, labelSize[1] + 10)

            cv2.rectangle(
                img_result,
                (x1, y1_label - labelSize[1] - 10),
                (x1 + labelSize[0], y1_label),
                color,
                -1
            )

            cv2.putText(
                img_result,
                label,
                (x1, y1_label - 5),
                cv2.FONT_HERSHEY_SIMPLEX,
                0.5,
                (255, 255, 255),
                1
            )

        # 반투명 적용
        img_result = cv2.addWeighted(overlay, 0.2, img_result, 0.8, 0)

        return cv2.cvtColor(img_result, cv2.COLOR_BGR2RGB)

    def create_text_visualization(self, image):
        """텍스트가 삽입된 문서 시각화"""
        canvas_height, canvas_width = image.shape[:2]
        canvas = np.ones((canvas_height, canvas_width, 3), dtype=np.uint8) * 255

        # OCR 및 API 바운딩 박스 그리기
        for result in self.ocr_results + self.api_results:
            x1, y1, x2, y2 = result['coordinates']
            cv2.rectangle(canvas, (x1, y1), (x2, y2), (0, 0, 0), 2)

        # PIL로 변환하여 한글 텍스트 추가
        canvas_pil = Image.fromarray(cv2.cvtColor(canvas, cv2.COLOR_BGR2RGB))
        draw = ImageDraw.Draw(canvas_pil)

        # 기본 폰트 사용 (Windows에서)
        try:
            font = ImageDraw.getfont()
        except:
            from PIL import ImageFont
            try:
                font = ImageFont.load_default()
            except:
                font = None

        # 텍스트 추가
        for result in self.ocr_results:
            x1, y1, x2, y2 = result['coordinates']
            text = result['text'].replace('\n', ' ')
            if len(text) > 50:
                text = text[:50] + "..."
            if font:
                draw.text((x1 + 5, y1 + 5), text, font=font, fill=(0, 0, 0))

        for result in self.api_results:
            x1, y1, x2, y2 = result['coordinates']
            text = result['description'].replace('\n', ' ')
            if len(text) > 50:
                text = text[:50] + "..."
            if font:
                draw.text((x1 + 5, y1 + 5), text, font=font, fill=(0, 0, 0))

        return np.array(canvas_pil)

    def create_cim_result(self, layout_info, ocr_results, ai_results):
        """CIM 결과 생성 (시각화 제거, JSON 통합만)"""
        from datetime import datetime
        
        # JSON 통합 결과 생성
        cim_result = {
            "document_structure": {
                "layout_analysis": {
                    "total_elements": len(layout_info),
                    "elements": []
                },
                "text_content": [],
                "ai_descriptions": []
            },
            "metadata": {
                "analysis_date": datetime.now().isoformat(),
                "total_text_regions": len([info for info in layout_info if 'text' in info.get('class_name', '').lower()]),
                "total_figures": len([info for info in layout_info if info.get('class_name') == 'figure']),
                "total_tables": len([info for info in layout_info if info.get('class_name') == 'table'])
            }
        }
        
        # 레이아웃 정보 통합
        for i, info in enumerate(layout_info):
            element = {
                "id": i,
                "class": info.get('class_name', 'unknown'),
                "confidence": float(info.get('confidence', 0.0)),
                "bbox": info.get('box', []),
                "area": info.get('area', 0)
            }
            
            # OCR 텍스트가 있는 경우 추가
            ocr_text = None
            for ocr_result in ocr_results:
                if ocr_result.get('id') == info.get('id'):
                    ocr_text = ocr_result.get('text', '')
                    break
            
            if ocr_text and ocr_text.strip():
                element["text"] = ocr_text
                cim_result["document_structure"]["text_content"].append({
                    "element_id": i,
                    "text": ocr_text,
                    "class": info.get('class_name', 'unknown')
                })
            
            # AI 설명이 있는 경우 추가
            ai_description = None
            for ai_result in ai_results:
                if ai_result.get('id') == info.get('id'):
                    ai_description = ai_result.get('description', '')
                    break
            
            if ai_description and ai_description.strip():
                element["ai_description"] = ai_description
                cim_result["document_structure"]["ai_descriptions"].append({
                    "element_id": i,
                    "description": ai_description,
                    "class": info.get('class_name', 'unknown')
                })
            
            cim_result["document_structure"]["layout_analysis"]["elements"].append(element)
        
        # 통계 계산
        text_elements = [e for e in cim_result["document_structure"]["layout_analysis"]["elements"] if "text" in e]
        ai_elements = [e for e in cim_result["document_structure"]["layout_analysis"]["elements"] if "ai_description" in e]
        
        stats = {
            "total_elements": len(layout_info),
            "text_elements": len(text_elements),
            "ai_described_elements": len(ai_elements),
            "class_distribution": {}
        }
        
        # 클래스별 분포 계산
        for info in layout_info:
            class_name = info.get('class_name', 'unknown')
            stats["class_distribution"][class_name] = stats["class_distribution"].get(class_name, 0) + 1
        
        return cim_result, stats


# 글로벌 분석기 인스턴스
analyzer = WorksheetAnalyzer()

# 구조화된 JSON 생성기 임포트
try:
    from structured_json_generator import StructuredJSONGenerator
    structured_generator = StructuredJSONGenerator()
    logger.info("✅ 구조화된 JSON 생성기가 성공적으로 로드되었습니다.")
except ImportError as e:
    logger.warning(f"⚠️ 구조화된 JSON 생성기 로드 실패: {e}")
    structured_generator = None


@app.post("/analyze-structured")
async def analyze_worksheet_structured(
    image: UploadFile = File(...),
    model_choice: str = Form("SmartEyeSsen"),
    api_key: Optional[str] = Form(None)
):
    """
    구조화된 학습지 분석 - 문제별 정렬 및 구조화
    """
    try:
        # 구조화된 JSON 생성기 확인
        if structured_generator is None:
            raise HTTPException(status_code=500, detail="구조화된 분석 기능을 사용할 수 없습니다.")
        
        # 이미지 읽기
        image_bytes = await image.read()
        pil_image = Image.open(io.BytesIO(image_bytes))
        
        # PIL 이미지를 OpenCV BGR 형태로 변환
        cv_image = cv2.cvtColor(np.array(pil_image), cv2.COLOR_RGB2BGR)
        
        # 모델 다운로드 및 로드
        logger.info(f"📊 구조화된 분석 시작 - 모델: {model_choice}")
        model_path = analyzer.download_model(model_choice)
        
        if not analyzer.load_model(model_path):
            raise HTTPException(status_code=500, detail="모델 로드에 실패했습니다.")
        
        # 레이아웃 분석
        layout_info = analyzer.analyze_layout(cv_image, model_choice)
        if not layout_info:
            raise HTTPException(status_code=400, detail="레이아웃 분석에 실패했습니다. 감지된 요소가 없습니다.")
        
        # OCR 처리
        analyzer.perform_ocr(cv_image)
        
        # OpenAI API 처리 (API 키가 있는 경우)
        if api_key and api_key.strip():
            analyzer.call_openai_api(cv_image, api_key)
        else:
            analyzer.api_results = []
        
        # 🆕 구조화된 JSON 생성
        logger.info("🔧 문제별 구조화 분석 수행 중...")
        structured_result = structured_generator.generate_structured_json(
            analyzer.ocr_results,
            analyzer.api_results,
            analyzer.layout_info
        )
        
        # 레이아웃 결과 시각화
        layout_viz = analyzer.visualize_results(cv_image)
        
        # 파일 저장
        timestamp = int(time.time())
        layout_viz_path = f"static/layout_viz_{timestamp}.png"
        
        layout_viz_pil = Image.fromarray(layout_viz)
        layout_viz_pil.save(layout_viz_path)
        
        # 구조화된 JSON 파일 저장
        from datetime import datetime
        structured_filename = f"structured_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        structured_filepath = f"static/{structured_filename}"
        
        with open(structured_filepath, 'w', encoding='utf-8') as f:
            json.dump(structured_result, f, indent=2, ensure_ascii=False)
        
        # 🆕 구조화된 텍스트 생성
        structured_text = create_structured_text(structured_result)
        
        # 통계 생성
        class_counts = Counter(item['class_name'] for item in layout_info)
        stats = {
            "total_layout_elements": len(layout_info),
            "ocr_text_blocks": len(analyzer.ocr_results),
            "ai_descriptions": len(analyzer.api_results),
            "total_questions": structured_result.get('document_info', {}).get('total_questions', 0),
            "class_counts": dict(class_counts)
        }
        
        logger.info(f"✅ 구조화된 분석 완료: {stats['total_questions']}개 문제 감지")
        
        return JSONResponse({
            "success": True,
            "layout_image_url": f"/{layout_viz_path}",
            "structured_json_url": f"/{structured_filepath}",
            "structured_result": structured_result,
            "structured_text": structured_text,
            "stats": stats,
            "timestamp": timestamp,
            "analysis_type": "structured"
        })
        
    except Exception as e:
        logger.error(f"구조화된 분석 중 오류 발생: {e}")
        raise HTTPException(status_code=500, detail=f"구조화된 분석 중 오류가 발생했습니다: {str(e)}")


@app.post("/analyze")
async def analyze_worksheet(
    image: UploadFile = File(...),
    model_choice: str = Form("SmartEyeSsen"),
    api_key: Optional[str] = Form(None),
    ocr_engine: str = Form("varco")  # 🆕 OCR 엔진 선택 추가
):
    """
    학습지 분석 메인 엔드포인트
    ocr_engine: "varco" 또는 "tesseract"
    """
    try:
        # 이미지 읽기
        image_bytes = await image.read()
        pil_image = Image.open(io.BytesIO(image_bytes))
        
        # PIL 이미지를 OpenCV BGR 형태로 변환
        cv_image = cv2.cvtColor(np.array(pil_image), cv2.COLOR_RGB2BGR)
        
        # 모델 다운로드 및 로드
        logger.info(f"모델 선택: {model_choice}")
        model_path = analyzer.download_model(model_choice)
        
        if not analyzer.load_model(model_path):
            raise HTTPException(status_code=500, detail="모델 로드에 실패했습니다.")
        
        # 레이아웃 분석
        layout_info = analyzer.analyze_layout(cv_image, model_choice)
        if not layout_info:
            raise HTTPException(status_code=400, detail="레이아웃 분석에 실패했습니다. 감지된 요소가 없습니다.")
        
        # OCR 처리 (엔진 선택)
        use_varco = ocr_engine.lower() == "varco"
        analyzer.perform_ocr(cv_image, use_varco=use_varco)
        
        # OpenAI API 처리 (API 키가 있는 경우)
        if api_key and api_key.strip():
            analyzer.call_openai_api(cv_image, api_key)
        else:
            analyzer.api_results = []
        
        # 레이아웃 결과 시각화
        layout_viz = analyzer.visualize_results(cv_image)
        
        # 레이아웃 결과 이미지를 파일로 저장
        timestamp = int(time.time())
        layout_viz_path = f"static/layout_viz_{timestamp}.png"
        
        layout_viz_pil = Image.fromarray(layout_viz)
        layout_viz_pil.save(layout_viz_path)
        
        # CIM 통합 결과 생성 (JSON 데이터만)
        cim_result, cim_stats = analyzer.create_cim_result(
            analyzer.layout_info, 
            analyzer.ocr_results, 
            analyzer.api_results
        )
        
        # 🆕 포맷팅된 텍스트 자동 생성
        formatted_text = create_formatted_text(cim_result)
        
        # JSON 파일 저장
        from datetime import datetime
        json_filename = f"analysis_result_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        json_filepath = f"static/{json_filename}"
        
        with open(json_filepath, 'w', encoding='utf-8') as f:
            json.dump(cim_result, f, indent=2, ensure_ascii=False)
        
        # 통계 생성
        class_counts = Counter(item['class_name'] for item in layout_info)
        stats = {
            "total_layout_elements": len(layout_info),
            "ocr_text_blocks": len(analyzer.ocr_results),
            "ai_descriptions": len(analyzer.api_results),
            "class_counts": dict(class_counts),
            "cim_stats": cim_stats
        }
        
        # OCR 텍스트 통합 (편집 가능한 형태로)
        combined_ocr_text = ""
        for result in analyzer.ocr_results:
            combined_ocr_text += f"[{result['class_name']}]\n{result['text']}\n\n"
        
        # AI 설명 통합
        combined_ai_text = ""
        for result in analyzer.api_results:
            combined_ai_text += f"[{result['class_name']}]\n{result['description']}\n\n"
        
        return JSONResponse({
            "success": True,
            "layout_image_url": f"/{layout_viz_path}",
            "json_url": f"/{json_filepath}",
            "stats": stats,
            "ocr_results": analyzer.ocr_results,
            "ai_results": analyzer.api_results,
            "ocr_text": combined_ocr_text.strip(),
            "ai_text": combined_ai_text.strip(),
            "formatted_text": formatted_text,  # 🆕 추가
            "timestamp": timestamp,
            "ocr_engine_used": ocr_engine  # 🆕 사용된 OCR 엔진 정보
        })
        
    except Exception as e:
        logger.error(f"분석 중 오류 발생: {e}")
        raise HTTPException(status_code=500, detail=f"분석 중 오류가 발생했습니다: {str(e)}")


@app.get("/")
async def root():
    """API 상태 확인"""
    return {
        "message": "SmartEyeSsen 학습지 분석 API",
        "version": "1.0.0",
        "device": device,
        "available_models": [
            "SmartEyeSsen",
            "docstructbench", 
            "doclaynet_docsynth",
            "docsynth300k"
        ]
    }


@app.get("/health")
async def health_check():
    """헬스 체크 엔드포인트"""
    return {"status": "healthy", "device": device}


@app.post("/format-text")
async def format_text_from_json(json_file: UploadFile = File(...)):
    """
    JSON 파일을 업로드받아서 포맷팅된 텍스트 생성
    """
    try:
        # JSON 파일 읽기
        json_content = await json_file.read()
        data = json.loads(json_content.decode('utf-8'))
        
        # 포맷팅된 텍스트 생성
        formatted_text = create_formatted_text(data)
        
        return JSONResponse({
            "success": True,
            "formatted_text": formatted_text,
            "message": "텍스트 포맷팅이 완료되었습니다."
        })
        
    except Exception as e:
        logger.error(f"텍스트 포맷팅 중 오류 발생: {e}")
        raise HTTPException(status_code=500, detail=f"텍스트 포맷팅 중 오류가 발생했습니다: {str(e)}")


@app.post("/save-as-word")
async def save_as_word(
    text: str = Form(...),
    filename: Optional[str] = Form("smarteye_document")
):
    """
    편집된 텍스트를 워드 문서로 저장
    """
    try:
        # 워드 문서 생성
        doc = Document()
        
        # 문서 제목 추가
        title = doc.add_heading('SmartEye OCR 분석 결과', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 현재 날짜 추가
        from datetime import datetime
        date_paragraph = doc.add_paragraph(f"생성일: {datetime.now().strftime('%Y년 %m월 %d일 %H시 %M분')}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # 구분선 추가
        doc.add_paragraph("=" * 50)
        
        # 텍스트 내용 추가 (줄바꿈 처리)
        lines = text.split('\n')
        current_paragraph = None
        
        for line in lines:
            line = line.strip()
            
            # 빈 줄 처리
            if not line:
                if current_paragraph is not None:
                    doc.add_paragraph()  # 빈 줄 추가
                continue
            
            # 제목 형태 감지 ([제목] 형식)
            if line.startswith('[') and line.endswith(']'):
                heading = doc.add_heading(line.strip('[]'), level=2)
                current_paragraph = None
            # 문제번호 형태 감지 (숫자. 형식)
            elif line.replace(' ', '').replace('.', '').isdigit() and '.' in line:
                heading = doc.add_heading(line, level=3)
                current_paragraph = None
            # 일반 텍스트
            else:
                paragraph = doc.add_paragraph(line)
                current_paragraph = paragraph
        
        # 파일명 정리 (확장자 제거)
        if filename.endswith('.docx'):
            filename = filename[:-5]
        
        # 타임스탬프 추가
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_filename = f"{filename}_{timestamp}.docx"
        
        # static 폴더에 저장
        filepath = f"static/{safe_filename}"
        doc.save(filepath)
        
        logger.info(f"워드 문서 저장 완료: {filepath}")
        
        return JSONResponse({
            "success": True,
            "message": "워드 문서가 성공적으로 생성되었습니다.",
            "filename": safe_filename,
            "download_url": f"/{filepath}",
            "timestamp": timestamp
        })
        
    except Exception as e:
        logger.error(f"워드 문서 생성 중 오류 발생: {e}")
        raise HTTPException(status_code=500, detail=f"워드 문서 생성 중 오류가 발생했습니다: {str(e)}")


@app.get("/download/{filename}")
async def download_file(filename: str):
    """
    생성된 파일 다운로드
    """
    try:
        filepath = f"static/{filename}"
        if not os.path.exists(filepath):
            raise HTTPException(status_code=404, detail="파일을 찾을 수 없습니다.")
        
        return FileResponse(
            path=filepath,
            filename=filename,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"파일 다운로드 중 오류 발생: {e}")
        raise HTTPException(status_code=500, detail=f"파일 다운로드 중 오류가 발생했습니다: {str(e)}")


def create_structured_text(structured_result):
    """
    구조화된 JSON 결과를 읽기 쉬운 텍스트로 변환
    """
    formatted_text = ""
    
    # 문서 정보
    doc_info = structured_result.get('document_info', {})
    formatted_text += f"📋 문서 분석 결과\n"
    formatted_text += f"총 문제 수: {doc_info.get('total_questions', 0)}개\n"
    formatted_text += f"레이아웃 유형: {doc_info.get('layout_type', '미확인')}\n\n"
    formatted_text += "=" * 50 + "\n\n"
    
    # 각 문제별 처리
    questions = structured_result.get('questions', [])
    
    for i, question in enumerate(questions, 1):
        question_num = question.get('question_number', f'문제{i}')
        section = question.get('section', '')
        
        # 문제 제목
        formatted_text += f"🔸 {question_num}"
        if section:
            formatted_text += f" ({section})"
        formatted_text += "\n\n"
        
        question_content = question.get('question_content', {})
        
        # 지문 (passage)
        passage = question_content.get('passage', '').strip()
        if passage:
            formatted_text += f"📖 지문:\n{passage}\n\n"
        
        # 주요 문제
        main_question = question_content.get('main_question', '').strip()
        if main_question:
            formatted_text += f"❓ 문제:\n{main_question}\n\n"
        
        # 선택지
        choices = question_content.get('choices', [])
        if choices:
            formatted_text += "📝 선택지:\n"
            for choice in choices:
                choice_num = choice.get('choice_number', '')
                choice_text = choice.get('choice_text', '')
                if choice_num and choice_text:
                    formatted_text += f"   {choice_num} {choice_text}\n"
                elif choice_text:
                    formatted_text += f"   • {choice_text}\n"
            formatted_text += "\n"
        
        # 이미지 설명
        images = question_content.get('images', [])
        if images:
            formatted_text += "🖼️ 이미지 설명:\n"
            for img in images:
                description = img.get('description', '').strip()
                if description:
                    formatted_text += f"   {description}\n"
            formatted_text += "\n"
        
        # 표 설명
        tables = question_content.get('tables', [])
        if tables:
            formatted_text += "📊 표 설명:\n"
            for table in tables:
                description = table.get('description', '').strip()
                if description:
                    formatted_text += f"   {description}\n"
            formatted_text += "\n"
        
        # 해설
        explanations = question_content.get('explanations', '').strip()
        if explanations:
            formatted_text += f"💡 해설:\n{explanations}\n\n"
        
        # AI 분석
        ai_analysis = question.get('ai_analysis', {})
        image_descriptions = ai_analysis.get('image_descriptions', [])
        table_analysis = ai_analysis.get('table_analysis', [])
        
        if image_descriptions or table_analysis:
            formatted_text += "🤖 AI 분석:\n"
            
            for img_desc in image_descriptions:
                desc = img_desc.get('description', '').strip()
                if desc:
                    formatted_text += f"   [이미지] {desc}\n"
            
            for table_desc in table_analysis:
                desc = table_desc.get('description', '').strip()
                if desc:
                    formatted_text += f"   [표] {desc}\n"
            
            formatted_text += "\n"
        
        # 문제 구분선
        if i < len(questions):
            formatted_text += "-" * 30 + "\n\n"
    
    return formatted_text.strip()


def create_formatted_text(json_data):
    """
    JSON 데이터를 기반으로 포맷팅된 텍스트 생성
    현재 클래스를 활용한 포맷팅 규칙 적용
    """
    
    # 포맷팅 규칙 정의
    formatting_rules = {
        'title': {
            'prefix': '',
            'suffix': '\n\n',  # 제목 후 두 줄 띄기
            'indent': 0
        },
        'question_number': {
            'prefix': '',
            'suffix': '. ',  # 문제번호 후 점과 공백
            'indent': 0
        },
        'question_type': {
            'prefix': '   ',  # 3칸 들여쓰기
            'suffix': '\n',
            'indent': 3
        },
        'question_text': {
            'prefix': '   ',  # 3칸 들여쓰기
            'suffix': '\n',
            'indent': 3
        },
        'plain_text': {
            'prefix': '',
            'suffix': '\n',
            'indent': 0
        },
        'table_caption': {
            'prefix': '\n',  # 표 제목 앞 한 줄 띄기
            'suffix': '\n',
            'indent': 0
        },
        'table_footnote': {
            'prefix': '',
            'suffix': '\n\n',  # 표 각주 후 두 줄 띄기
            'indent': 0
        },
        'isolated_formula': {
            'prefix': '\n',  # 수식 앞 한 줄 띄기
            'suffix': '\n\n',  # 수식 후 두 줄 띄기
            'indent': 0
        },
        'formula_caption': {
            'prefix': '',
            'suffix': '\n',
            'indent': 0
        },
        'abandon_text': {
            'prefix': '[삭제됨] ',
            'suffix': '\n',
            'indent': 0
        },
        'figure': {
            'prefix': '\n[그림 설명] ',  # 그림 앞 한 줄 띄기
            'suffix': '\n\n',  # 그림 후 두 줄 띄기
            'indent': 0
        },
        'table': {
            'prefix': '\n[표 설명] ',  # 표 앞 한 줄 띄기
            'suffix': '\n\n',  # 표 후 두 줄 띄기
            'indent': 0
        }
    }
    
    # 요소들을 위치 기준으로 정렬 (y 좌표 기준)
    elements = []
    
    # 레이아웃 분석 요소에서 텍스트/AI 설명 추출
    if "document_structure" in json_data:
        layout_elements = json_data["document_structure"]["layout_analysis"]["elements"]
        
        for element in layout_elements:
            element_id = element.get("id")
            class_name = element.get("class", "").lower().replace(" ", "_")
            bbox = element.get("bbox", [0, 0, 0, 0])
            
            # OCR 텍스트 또는 AI 설명 찾기
            content = None
            content_type = None
            
            # OCR 텍스트 확인
            if "text" in element:
                content = element["text"]
                content_type = "ocr"
            # AI 설명 확인
            elif "ai_description" in element:
                content = element["ai_description"]
                content_type = "ai"
            
            if content and content.strip():
                elements.append({
                    'id': element_id,
                    'class': class_name,
                    'content': content.strip(),
                    'type': content_type,
                    'y_position': bbox[1] if len(bbox) > 1 else 0,  # y 좌표
                    'x_position': bbox[0] if len(bbox) > 0 else 0   # x 좌표
                })
    
    # Y 좌표 기준으로 정렬 (위에서 아래로)
    elements.sort(key=lambda x: (x['y_position'], x['x_position']))
    
    # 포맷팅된 텍스트 생성
    formatted_text = ""
    prev_class = None
    
    for element in elements:
        class_name = element['class']
        content = element['content']
        
        # 포맷팅 규칙 적용
        rule = formatting_rules.get(class_name, {
            'prefix': '',
            'suffix': '\n',
            'indent': 0
        })
        
        # 특별한 조건 처리
        formatted_line = ""
        
        # 문제번호와 문제텍스트가 연속으로 나오는 경우 처리
        if class_name == 'question_text' and prev_class == 'question_number':
            # 문제번호 바로 뒤에 문제텍스트가 오면 같은 줄에 배치
            formatted_line = content + rule['suffix']
        else:
            # 일반적인 포맷팅 적용
            prefix = rule['prefix']
            suffix = rule['suffix']
            
            formatted_line = prefix + content + suffix
        
        formatted_text += formatted_line
        prev_class = class_name
    
    # 최종 정리 (연속된 빈 줄 정리)
    lines = formatted_text.split('\n')
    cleaned_lines = []
    prev_empty = False
    
    for line in lines:
        is_empty = line.strip() == ''
        
        # 연속된 빈 줄이 3개 이상 나오지 않도록 제한
        if is_empty and prev_empty:
            continue
        
        cleaned_lines.append(line)
        prev_empty = is_empty
    
    return '\n'.join(cleaned_lines).strip()


@app.post("/compare-ocr")
async def compare_ocr_engines(
    image: UploadFile = File(...),
    model_choice: str = Form("SmartEyeSsen")
):
    """
    Tesseract와 VARCO Vision OCR 성능 비교
    """
    try:
        # 이미지 읽기
        image_bytes = await image.read()
        pil_image = Image.open(io.BytesIO(image_bytes))
        cv_image = cv2.cvtColor(np.array(pil_image), cv2.COLOR_RGB2BGR)
        
        # 모델 로드
        model_path = analyzer.download_model(model_choice)
        if not analyzer.load_model(model_path):
            raise HTTPException(status_code=500, detail="모델 로드에 실패했습니다.")
        
        # 레이아웃 분석
        layout_info = analyzer.analyze_layout(cv_image, model_choice)
        
        # Tesseract OCR
        start_time = time.time()
        tesseract_results = analyzer.perform_ocr(cv_image, use_varco=False)
        tesseract_time = time.time() - start_time
        
        # VARCO Vision OCR
        start_time = time.time()
        varco_results = analyzer.perform_ocr(cv_image, use_varco=True)
        varco_time = time.time() - start_time
        
        # 결과 비교 분석
        comparison = {
            "tesseract": {
                "processing_time": round(tesseract_time, 2),
                "text_blocks": len(tesseract_results),
                "total_characters": sum(len(r['text']) for r in tesseract_results),
                "results": tesseract_results
            },
            "varco": {
                "processing_time": round(varco_time, 2),
                "text_blocks": len(varco_results),
                "total_characters": sum(len(r['text']) for r in varco_results),
                "results": varco_results
            },
            "comparison": {
                "speed_ratio": round(tesseract_time / varco_time if varco_time > 0 else 0, 2),
                "accuracy_note": "정확도 비교를 위해서는 실제 정답 데이터가 필요합니다."
            }
        }
        
        return JSONResponse({
            "success": True,
            "comparison": comparison,
            "message": "OCR 엔진 비교 완료"
        })
        
    except Exception as e:
        logger.error(f"OCR 비교 중 오류 발생: {e}")
        raise HTTPException(status_code=500, detail=f"OCR 비교 중 오류가 발생했습니다: {str(e)}")


if __name__ == "__main__":
    print("🚀 SmartEyeSsen API 서버를 시작합니다...")
    print(f"📱 브라우저에서 http://localhost:8000 으로 접속하세요")
    print(f"📚 API 문서는 http://localhost:8000/docs 에서 확인할 수 있습니다")
    
    uvicorn.run(
        "api_server:app",
        host="0.0.0.0",
        port=8000,
        reload=True,
        log_level="info"
    )
